"""
Liberty In a Can - SEO Content Agent
Runs daily via GitHub Actions. Saves .docx files to Google Drive (opens as Google Docs).
"""
import anthropic
import os
import io
import json
import re
import time
from datetime import datetime
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from docx import Document
from docx.shared import Pt

ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]
GOOGLE_DRIVE_FOLDER_ID = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "")
GOOGLE_SERVICE_ACCOUNT_JSON = os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"]

KEYWORD_POOL = [
    "THC seltzer Florida",
    "hemp drinks alcohol alternative",
    "THC beverages near me",
    "cannabis seltzer vs alcohol",
    "hemp THC drink review",
    "best THC drinks 2025",
    "sober curious drinks with THC",
    "THC tea Florida",
    "low dose THC drinks",
    "hemp infused seltzer",
    "alcohol alternative THC beverage",
    "social THC drink",
    "THC sparkling water",
    "hemp drink Tennessee",
    "cannabis beverage distribution",
]

BRAND_VOICE = """You are the official SEO copywriter for Liberty In a Can.
Liberty In a Can makes hemp-derived THC beverages: Liberty Seltzer and Liberty Tea (5mg and 10mg).
Distributed in Florida and Tennessee. Brand voice: American freedom, bold, witty, lifestyle-forward.
Never make health claims. Audience: adults 25-55, sober-curious, reducing alcohol.
Use markdown formatting with # for H1, ## for H2, ** for bold, * for bullets.
Lead with primary keyword naturally in first 100 words."""


def pick_todays_keyword():
    override = os.environ.get("KEYWORD_OVERRIDE", "").strip()
    if override:
        return override
    day_of_year = datetime.now().timetuple().tm_yday
    return KEYWORD_POOL[day_of_year % len(KEYWORD_POOL)]


def call_with_retry(fn, max_retries=3):
    for attempt in range(max_retries):
        try:
            return fn()
        except anthropic.RateLimitError as e:
            wait = 60 * (attempt + 1)
            print(f"Rate limit hit, waiting {wait}s before retry {attempt + 1}/{max_retries}...")
            time.sleep(wait)
    raise Exception("Max retries exceeded")


def research_keyword(client, keyword):
    print(f"Researching: {keyword}")
    def _call():
        return client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            system="You are an SEO research specialist for hemp THC beverages. Search for the keyword and return a brief with: search intent, top 5 LSI keywords, what top content covers, and content hooks. Be concise.",
            messages=[{"role": "user", "content": f'Research SEO landscape for: "{keyword}". Focus on hemp THC beverages and alcohol alternatives.'}]
        )
    response = call_with_retry(_call)
    research = "\n".join(block.text for block in response.content if block.type == "text")
    print(f"Research complete ({len(research.split())} words)")
    return research


def write_blog_post(client, keyword, research):
    print("Writing blog post...")
    time.sleep(30)
    def _call():
        return client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1500,
            system=BRAND_VOICE,
            messages=[{"role": "user", "content": f"""Write a full SEO blog post (800-1000 words) targeting: "{keyword}"
Structure: # H1 title, ## H2 sections (3-4), conclusion, ## Frequently Asked Questions (3 Q&As)
SEO Research: {research}
Use 2-3 related keywords naturally. Use markdown formatting throughout."""}]
        )
    response = call_with_retry(_call)
    content = response.content[0].text
    print(f"Blog post written ({len(content.split())} words)")
    return content


def write_product_copy(client, keyword, research):
    print("Writing product copy...")
    time.sleep(30)
    def _call():
        return client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=600,
            system=BRAND_VOICE,
            messages=[{"role": "user", "content": f"""Write SEO product page copy targeting: "{keyword}"
Structure: # H1 headline, ## Benefits with 3 bullet points, ## Description (120-150 words), ## Call to Action, **Meta Title:** (max 60 chars), **Meta Description:** (max 160 chars)
SEO Research: {research}"""}]
        )
    response = call_with_retry(_call)
    content = response.content[0].text
    print(f"Product copy written ({len(content.split())} words)")
    return content


def markdown_to_docx(markdown_text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for line in markdown_text.split('\n'):
        line = line.rstrip()
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph()
            for part in re.split(r'(\*\*.*?\*\*)', line):
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    return doc


def get_drive_service():
    creds_dict = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)
    creds = service_account.Credentials.from_service_account_info(
        creds_dict, scopes=["https://www.googleapis.com/auth/drive"]
    )
    return build("drive", "v3", credentials=creds)


def save_docx_to_drive(drive_service, filename, doc, folder_id=None):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    file_metadata = {
        "name": filename,
        "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    if folder_id:
        file_metadata["parents"] = [folder_id]
    media = MediaIoBaseUpload(buffer, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    file = drive_service.files().create(body=file_metadata, media_body=media, fields="id").execute()
    url = f"https://drive.google.com/file/d/{file.get('id')}/view"
    print(f"Saved: {filename} -> {url}")
    return url


def run_agent():
    print("Liberty In a Can - SEO Agent Starting")
    print(datetime.now().strftime("%A, %B %d %Y at %I:%M %p"))
    print("-" * 50)

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    keyword = pick_todays_keyword()
    print(f"Today's keyword: {keyword}")

    research = research_keyword(client, keyword)
    blog_post = write_blog_post(client, keyword, research)
    product_copy = write_product_copy(client, keyword, research)

    print("\nSaving to Google Drive...")
    drive_service = get_drive_service()
    date_str = datetime.now().strftime("%Y-%m-%d")
    safe_kw = keyword.replace(" ", "_").replace("/", "-")
    folder_id = GOOGLE_DRIVE_FOLDER_ID or None

    blog_doc = markdown_to_docx(f"**Keyword:** {keyword}\n**Date:** {date_str}\n\n---\n\n{blog_post}")
    save_docx_to_drive(drive_service, f"LIAC_Blog_{safe_kw}_{date_str}.docx", blog_doc, folder_id)

    product_doc = markdown_to_docx(f"**Keyword:** {keyword}\n**Date:** {date_str}\n\n---\n\n{product_copy}")
    save_docx_to_drive(drive_service, f"LIAC_ProductCopy_{safe_kw}_{date_str}.docx", product_doc, folder_id)

    print("\nDone! Check your Google Drive LIAC SEO Content folder.")


if __name__ == "__main__":
    run_agent()
